import os
import uuid
from typing import List, Dict, Any
from PyPDF2 import PdfReader
from docx import Document
import markdown
import re
import traceback
from pdf2image import convert_from_path, convert_from_bytes
import pytesseract
from io import BytesIO

class DocumentProcessor:
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.md', '.txt']
        self.chunk_size = 2000  # Verhoogd van 1000 naar 2000
        self.chunk_overlap = 200  # Overlap tussen chunks voor betere context
    
    def process_document(self, file_path: str, file_content: bytes = None) -> List[str]:
        """Verwerk een document en splits het in chunks"""
        try:
            # Extract text based on file type
            text = self._extract_text(file_path, file_content)
            if not text:
                return []
            
            # Clean and normalize text
            text = self._clean_text(text)
            
            # Split into chunks with better handling for invoices
            chunks = self._split_into_chunks(text)
            
            print(f"Processed {file_path}: {len(chunks)} chunks created")
            return chunks
        except Exception as e:
            print(f"Error processing document {file_path}: {e}")
            return []
    
    def _clean_text(self, text: str) -> str:
        """Clean en normaliseer tekst voor betere verwerking"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Fix common OCR issues in invoices
        text = re.sub(r'(\d+)\s*,\s*(\d{2})', r'\1,\2', text)  # Fix decimal numbers
        text = re.sub(r'(\d+)\s*\.\s*(\d{2})', r'\1,\2', text)  # Fix decimal numbers with dots
        
        # Normalize common invoice terms
        replacements = {
            'betalen': 'betalen',
            'betaling': 'betaling',
            'voldoen': 'voldoen',
            'voldening': 'voldening',
            'factuur': 'factuur',
            'rekening': 'rekening',
            'bedrag': 'bedrag',
            'totaal': 'totaal',
            'euro': 'euro',
            '€': 'euro',
            'EUR': 'euro'
        }
        
        for old, new in replacements.items():
            text = re.sub(rf'\b{old}\b', new, text, flags=re.IGNORECASE)
        
        return text.strip()
    
    def _split_into_chunks(self, text: str) -> List[str]:
        """Split tekst in chunks met betere logica voor facturen"""
        if not text:
            return []
        
        # Voor facturen: probeer eerst op natuurlijke grenzen te splitsen
        chunks = []
        
        # Split op paragrafen (dubbele newlines)
        paragraphs = re.split(r'\n\s*\n', text)
        
        current_chunk = ""
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # Als deze paragraaf de chunk te groot maakt, start een nieuwe
            if len(current_chunk) + len(paragraph) > 3000:  # Grotere chunks voor facturen
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = paragraph
            else:
                if current_chunk:
                    current_chunk += "\n\n" + paragraph
                else:
                    current_chunk = paragraph
        
        # Voeg laatste chunk toe
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Als we te weinig chunks hebben, split verder op zinnen
        if len(chunks) < 2:
            chunks = []
            sentences = re.split(r'[.!?]+', text)
            
            current_chunk = ""
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                
                if len(current_chunk) + len(sentence) > 3000:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence
                else:
                    if current_chunk:
                        current_chunk += ". " + sentence
                    else:
                        current_chunk = sentence
            
            if current_chunk:
                chunks.append(current_chunk.strip())
        
        # Als laatste redmiddel: fixed-size chunks
        if not chunks:
            chunks = [text[i:i+3000] for i in range(0, len(text), 2500)]  # 500 overlap
        
        # Clean chunks
        cleaned_chunks = []
        for chunk in chunks:
            chunk = chunk.strip()
            if len(chunk) > 50:  # Minimum chunk size
                cleaned_chunks.append(chunk)
        
        return cleaned_chunks 
    
    def _extract_text(self, file_path: str, file_content: bytes = None) -> str:
        """Extract tekst uit verschillende bestandstypen"""
        try:
            file_type = self._get_file_type(file_path)
            
            if file_type == 'pdf':
                return self._extract_pdf_text(file_path, file_content)
            elif file_type == 'docx':
                return self._extract_docx_text(file_path, file_content)
            elif file_type in ['txt', 'md']:
                return self._extract_text_file(file_path, file_content)
            else:
                print(f"Unsupported file type: {file_type}")
                return ""
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _get_file_type(self, file_path: str) -> str:
        """Bepaal bestandstype op basis van extensie"""
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            return 'docx'
        elif ext == '.txt':
            return 'txt'
        elif ext == '.md':
            return 'md'
        else:
            return 'unknown'
    
    def _extract_pdf_text(self, file_path: str, file_content: bytes = None) -> str:
        """Extract tekst uit PDF met OCR fallback"""
        try:
            # Use file_content if provided, otherwise read from file
            if file_content:
                pdf_reader = PdfReader(BytesIO(file_content))
            else:
                # Read file content first to avoid file handle issues
                with open(file_path, 'rb') as file:
                    file_content = file.read()
                pdf_reader = PdfReader(BytesIO(file_content))
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if not text or not text.strip():
                    # OCR fallback
                    print(f"Page {page_num+1}: no text found, trying OCR...")
                    try:
                        if file_content:
                            images = convert_from_bytes(file_content)
                        else:
                            images = convert_from_path(file_path)
                        
                        if page_num < len(images):
                            ocr_text = pytesseract.image_to_string(images[page_num], lang='nld+eng')
                            text = ocr_text
                            print(f"Page {page_num+1} OCR successful: {len(text)} characters")
                        else:
                            print(f"Page {page_num+1}: no image available for OCR")
                    except Exception as ocr_error:
                        print(f"OCR failed for page {page_num+1}: {ocr_error}")
                        text = ""
                
                if text and text.strip():
                    text_parts.append(text.strip())
                    print(f"Page {page_num+1}: extracted {len(text)} characters")
                else:
                    print(f"Page {page_num+1}: no text extracted")
            
            result = '\n\n'.join(text_parts)
            print(f"Total extracted text: {len(result)} characters")
            return result
        except Exception as e:
            print(f"Error extracting PDF text: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: str, file_content: bytes = None) -> str:
        """Extract tekst uit DOCX bestand"""
        try:
            if file_content:
                doc = Document(BytesIO(file_content))
            else:
                doc = Document(file_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            return '\n\n'.join(text_parts)
        except Exception as e:
            print(f"Error extracting DOCX text: {e}")
            return ""
    
    def _extract_text_file(self, file_path: str, file_content: bytes = None) -> str:
        """Extract tekst uit tekst bestand"""
        try:
            if file_content:
                content = file_content.decode('utf-8')
            else:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
            
            # Convert markdown to plain text if needed
            if file_path.endswith('.md'):
                content = markdown.markdown(content)
                # Remove HTML tags
                content = re.sub(r'<[^>]+>', '', content)
            
            return content
        except Exception as e:
            print(f"Error extracting text file: {e}")
            return "" 